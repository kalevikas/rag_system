"""
Multi-Format Document Processor
Supports: PDF, Excel (.xlsx/.xls), CSV, Word (.docx/.doc),
          Plain Text (.txt), Markdown (.md), URLs, REST APIs
"""
import os
import io
import csv
import json
import logging
import tempfile
import hashlib
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional, Callable

from langchain_core.documents import Document

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def _clean_text(text: str) -> str:
    """Remove excessive whitespace while preserving structure."""
    import re
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    text = re.sub(r' {3,}', '  ', text)
    return text.strip()


def _build_metadata(source: str, doc_type: str, extra: Optional[Dict] = None) -> Dict:
    meta = {
        "source": source,
        "doc_type": doc_type,
        "processed_date": datetime.utcnow().isoformat(),
    }
    if extra:
        meta.update(extra)
    return meta


# ---------------------------------------------------------------------------
# PDF Processor  (wraps existing PDFProcessor for consistency)
# ---------------------------------------------------------------------------

def load_pdf(file_path: str) -> List[Document]:
    """Load a PDF file and return LangChain Documents (one per page)."""
    docs: List[Document] = []
    try:
        import fitz  # PyMuPDF
        pdf = fitz.open(file_path)
        for page_num, page in enumerate(pdf, start=1):
            text = page.get_text("text")
            if text.strip():
                docs.append(Document(
                    page_content=_clean_text(text),
                    metadata=_build_metadata(
                        file_path, "pdf",
                        {"file_name": os.path.basename(file_path), "page_number": page_num}
                    )
                ))
        pdf.close()
        logger.info(f"[PDF] Loaded {len(docs)} pages from {file_path}")
    except Exception as e:
        logger.error(f"[PDF] Failed to load {file_path}: {e}")
        # Fallback: pypdf
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            for i, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    docs.append(Document(
                        page_content=_clean_text(text),
                        metadata=_build_metadata(
                            file_path, "pdf",
                            {"file_name": os.path.basename(file_path), "page_number": i}
                        )
                    ))
        except Exception as e2:
            logger.error(f"[PDF] Fallback pypdf also failed: {e2}")
    return docs


# ---------------------------------------------------------------------------
# Excel / CSV Processor
# ---------------------------------------------------------------------------

def _dataframe_to_docs(df, source: str, sheet_name: str = "Sheet1") -> List[Document]:
    """Convert a pandas DataFrame to LangChain Documents row-by-row and as summary."""
    docs: List[Document] = []
    # Full table summary document
    summary_text = f"Table: {sheet_name}\nColumns: {', '.join(str(c) for c in df.columns)}\nRows: {len(df)}\n\n"
    summary_text += df.to_string(index=False, max_rows=500)
    docs.append(Document(
        page_content=_clean_text(summary_text),
        metadata=_build_metadata(source, "table", {"sheet": sheet_name, "rows": len(df)})
    ))
    # Each row as a separate document for fine-grained retrieval
    for idx, row in df.iterrows():
        row_text = " | ".join(f"{col}: {val}" for col, val in row.items() if str(val).strip() not in ("", "nan"))
        if row_text.strip():
            docs.append(Document(
                page_content=row_text,
                metadata=_build_metadata(source, "row", {"sheet": sheet_name, "row_index": idx})
            ))
    return docs


def load_excel(file_path: str) -> List[Document]:
    """Load an Excel file (.xlsx / .xls) — all sheets."""
    docs: List[Document] = []
    try:
        import pandas as pd
        xl = pd.ExcelFile(file_path)
        for sheet in xl.sheet_names:
            df = xl.parse(sheet)
            df = df.fillna("").astype(str)
            docs.extend(_dataframe_to_docs(df, file_path, sheet))
        logger.info(f"[Excel] Loaded {len(docs)} documents from {file_path} ({len(xl.sheet_names)} sheets)")
    except Exception as e:
        logger.error(f"[Excel] Failed: {e}")
    return docs


def load_csv(file_path: str) -> List[Document]:
    """Load a CSV file."""
    docs: List[Document] = []
    try:
        import pandas as pd
        df = pd.read_csv(file_path, encoding="utf-8", on_bad_lines="skip")
        df = df.fillna("").astype(str)
        docs.extend(_dataframe_to_docs(df, file_path, Path(file_path).stem))
        logger.info(f"[CSV] Loaded {len(docs)} documents from {file_path}")
    except Exception as e:
        logger.error(f"[CSV] Failed: {e}")
    return docs


# ---------------------------------------------------------------------------
# Word Document Processor (.docx / .doc)
# ---------------------------------------------------------------------------

def load_word(file_path: str) -> List[Document]:
    """Load a Word document (.docx)."""
    docs: List[Document] = []
    try:
        import docx
        doc = docx.Document(file_path)
        full_text = []
        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text.append(row_text)
        combined = "\n\n".join(full_text)
        if combined.strip():
            docs.append(Document(
                page_content=_clean_text(combined),
                metadata=_build_metadata(file_path, "docx",
                                          {"file_name": os.path.basename(file_path)})
            ))
        logger.info(f"[Word] Loaded {file_path}")
    except Exception as e:
        logger.error(f"[Word] Failed: {e}")
    return docs


# ---------------------------------------------------------------------------
# Plain Text / Markdown Processor
# ---------------------------------------------------------------------------

def load_text(file_path: str) -> List[Document]:
    """Load a plain text or Markdown file."""
    docs: List[Document] = []
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        ext = Path(file_path).suffix.lower()
        doc_type = "markdown" if ext in (".md", ".markdown") else "text"
        if text.strip():
            docs.append(Document(
                page_content=_clean_text(text),
                metadata=_build_metadata(file_path, doc_type,
                                          {"file_name": os.path.basename(file_path)})
            ))
        logger.info(f"[Text] Loaded {file_path}")
    except Exception as e:
        logger.error(f"[Text] Failed: {e}")
    return docs


# ---------------------------------------------------------------------------
# REST API Processor
# ---------------------------------------------------------------------------

def load_api(
    url: str,
    method: str = "GET",
    headers: Optional[Dict] = None,
    params: Optional[Dict] = None,
    body: Optional[Dict] = None,
    json_path: Optional[str] = None,   # e.g. "results.items" to drill into nested JSON
) -> List[Document]:
    """
    Fetch data from a REST API endpoint and convert to Documents.

    Args:
        url: API endpoint URL
        method: HTTP method (GET / POST)
        headers: Optional HTTP headers
        params: Optional query parameters
        body: Optional request body (for POST)
        json_path: Dot-separated path to extract from JSON response
    """
    docs: List[Document] = []
    try:
        import requests
        resp = requests.request(
            method.upper(),
            url,
            headers=headers or {},
            params=params or {},
            json=body,
            timeout=30
        )
        resp.raise_for_status()
        data = resp.json()

        # Drill into json_path if specified
        if json_path:
            for key in json_path.split("."):
                data = data[key]

        # Convert to text
        if isinstance(data, list):
            for i, item in enumerate(data):
                text = json.dumps(item, ensure_ascii=False, indent=2) if isinstance(item, dict) else str(item)
                docs.append(Document(
                    page_content=_clean_text(text),
                    metadata=_build_metadata(url, "api", {"item_index": i})
                ))
        elif isinstance(data, dict):
            docs.append(Document(
                page_content=_clean_text(json.dumps(data, ensure_ascii=False, indent=2)),
                metadata=_build_metadata(url, "api")
            ))
        else:
            docs.append(Document(
                page_content=_clean_text(str(data)),
                metadata=_build_metadata(url, "api")
            ))
        logger.info(f"[API] Loaded {len(docs)} records from {url}")
    except Exception as e:
        logger.error(f"[API] Failed: {e}")
    return docs


# ---------------------------------------------------------------------------
# URL / Web Scraper (lightweight BeautifulSoup + Selenium fallback)
# ---------------------------------------------------------------------------

def load_url(
    url: str,
    use_selenium: bool = False,
    depth: int = 1,
    callback: Optional[Callable[[Document], None]] = None
) -> List[Document]:
    """
    Scrape a URL. If use_selenium=True uses the existing WebScraper,
    otherwise uses requests + BeautifulSoup.
    """
    docs: List[Document] = []
    try:
        if use_selenium:
            from src.web_scraper import WebScraper
            scraper = WebScraper(headless=True, follow_links=(depth > 1), max_depth=depth)
            scraped = scraper.scrape_urls([url], max_pages=30, callback=callback)
            return scraped or []

        import requests
        from bs4 import BeautifulSoup
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 Chrome/120 Safari/537.36"
            )
        }
        resp = requests.get(url, headers=headers, timeout=20)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")

        # Remove noise
        for tag in soup(["script", "style", "nav", "footer", "header", "form", "iframe"]):
            tag.decompose()

        title = soup.title.string.strip() if soup.title else url
        text = soup.get_text(separator="\n", strip=True)
        import re
        text = re.sub(r'\n{3,}', '\n\n', text)

        if text.strip():
            docs.append(Document(
                page_content=_clean_text(text),
                metadata=_build_metadata(url, "url", {"title": title})
            ))
        logger.info(f"[URL] Loaded {url}")
    except Exception as e:
        logger.error(f"[URL] Failed to load {url}: {e}")
    return docs


# ---------------------------------------------------------------------------
# Universal Dispatcher
# ---------------------------------------------------------------------------

EXTENSION_MAP = {
    ".pdf":  load_pdf,
    ".xlsx": load_excel,
    ".xls":  load_excel,
    ".csv":  load_csv,
    ".docx": load_word,
    ".doc":  load_word,
    ".txt":  load_text,
    ".md":   load_text,
    ".markdown": load_text,
}


def load_file(file_path: str) -> List[Document]:
    """
    Automatically detect file type and load it into LangChain Documents.
    Supports: PDF, Excel, CSV, Word, Text, Markdown.
    """
    ext = Path(file_path).suffix.lower()
    loader_fn = EXTENSION_MAP.get(ext)
    if loader_fn is None:
        logger.warning(f"[MultiFormat] Unknown extension '{ext}' — treating as plain text")
        return load_text(file_path)
    return loader_fn(file_path)


def get_supported_extensions() -> List[str]:
    """Return list of supported file extensions."""
    return list(EXTENSION_MAP.keys())
